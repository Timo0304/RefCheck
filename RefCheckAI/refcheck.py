import requests
from fuzzywuzzy import fuzz
from Bio import Entrez
from docx import Document
import fitz
import streamlit as st
import openai
import pandas as pd
from fpdf import FPDF
import re

# Configure your credentials
Entrez.email = "your-email@example.com"  # Replace with your email
openai.api_key = "your-openai-api-key"  # Replace with your OpenAI API key

# Utility function to clean text for DOCX and other formats
def clean_text(text):
    if not isinstance(text, str):
        return str(text)
    return re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)

# Search & Suggestion Functions
def search_crossref(title):
    url = "https://api.crossref.org/works"
    params = {"query.title": title, "rows": 3}
    response = requests.get(url, params=params)
    if response.status_code == 200:
        return response.json()["message"].get("items", [])
    return []

def search_pubmed(reference):
    try:
        handle = Entrez.esearch(db="pubmed", term=reference, retmax=1)
        record = Entrez.read(handle)
        if record["IdList"]:
            pubmed_id = record["IdList"][0]
            summary_handle = Entrez.esummary(db="pubmed", id=pubmed_id)
            summary = Entrez.read(summary_handle)
            doc = summary[0]
            return {
                "title": doc.get("Title", ""),
                "doi": doc.get("DOI", ""),
                "authors": doc.get("Source", ""),
                "year": doc.get("PubDate", "")
            }
    except Exception as e:
        print(f"PubMed error: {e}")
    return None

def search_semantic_scholar(reference):
    try:
        url = f"https://api.semanticscholar.org/graph/v1/paper/search?query={reference}&limit=1&fields=title,authors,year,venue,doi"
        res = requests.get(url)
        data = res.json()
        if data.get("data"):
            paper = data["data"][0]
            return {
                "title": paper.get("title"),
                "doi": paper.get("doi"),
                "authors": ", ".join([a['name'] for a in paper.get("authors", [])]),
                "year": paper.get("year"),
                "journal": paper.get("venue")
            }
    except Exception as e:
        print(f"Semantic Scholar error: {e}")
    return None

def suggest_correction(reference, style):
    prompt = f"Given the reference below, correct it and reformat it to proper {style} citation style:\n\n{reference}"
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message['content'].strip()
    except Exception as e:
        return f"[GPT error: {e}]"

# Main Verification Logic
def verify_reference(reference, style="APA"):
    result = {
        "Input": reference,
        "Matched": "",
        "DOI": "",
        "Status": "",
        "Correction": ""
    }
    candidates = search_crossref(reference)
    best_match = None
    best_score = 0

    for item in candidates:
        candidate_title = item.get("title", [""])[0]
        score = fuzz.token_set_ratio(reference.lower(), candidate_title.lower())
        if score > best_score:
            best_score = score
            best_match = item

    if best_match and best_score >= 60:
        result["Matched"] = best_match.get("title", [""])[0]
        result["DOI"] = best_match.get("DOI", "")
        result["Status"] = "✅ VERIFIED" if best_score >= 80 else "⚠️ POSSIBLY INCORRECT"
    else:
        fallback = search_pubmed(reference) or search_semantic_scholar(reference)
        if fallback:
            result["Matched"] = fallback["title"]
            result["DOI"] = fallback.get("doi", "")
            result["Status"] = "⚠️ POSSIBLY CORRECT (via fallback)"
        else:
            result["Status"] = "❌ No match"
            result["Correction"] = suggest_correction(reference, style)
    return result

# Extractors
def extract_references_from_pdf(file_path):
    doc = fitz.open(file_path)
    references = []
    for page in doc:
        text = page.get_text()
        if "References" in text:
            ref_section = text.split("References", 1)[-1]
            lines = ref_section.split("\n")
            for line in lines:
                if line.strip() and any(char.isdigit() for char in line):
                    references.append(line.strip())
            break
    return references

def extract_references_from_docx(file_path):
    doc = Document(file_path)
    references = []
    start_collecting = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if not start_collecting and "references" in text.lower():
            start_collecting = True
            continue
        if start_collecting and text:
            references.append(text)
    return references

# Exporters
def export_to_docx(df, filename="verified_references.docx"):
    doc = Document()
    doc.add_heading("Verified References", level=1)
    for i, row in df.iterrows():
        doc.add_paragraph(
            f"{i + 1}. Input: {clean_text(row['Input'])}\n"
            f"Matched: {clean_text(row['Matched'])}\n"
            f"DOI: {clean_text(row['DOI'])}\n"
            f"Status: {clean_text(row['Status'])}\n"
            f"Correction: {clean_text(row['Correction'])}"
        )
    doc.save(filename)

def export_to_pdf(df, filename="verified_references.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=10)
    pdf.cell(200, 10, txt="Verified References", ln=True, align="C")
    pdf.ln(5)
    for i, row in df.iterrows():
        text = f"{i + 1}. Input: {clean_text(row['Input'])}\nMatched: {clean_text(row['Matched'])}\nDOI: {clean_text(row['DOI'])}\nStatus: {clean_text(row['Status'])}\nCorrection: {clean_text(row['Correction'])}\n"
        pdf.multi_cell(0, 10, text)
        pdf.ln(1)
    pdf.output(filename)

# Streamlit UI
def streamlit_interface():
    st.title("📚 RefCheckAI: Reference Verifier")
    style = st.selectbox("Select Reference Style", ["APA", "Vancouver", "MLA", "Chicago"])
    uploaded_file = st.file_uploader("Upload a PDF, DOCX or TXT file", type=["pdf", "docx", "txt"])
    manual_refs = st.text_area("Or paste your references (one per line)")

    if st.button("Verify References"):
        references = []
        if uploaded_file:
            if uploaded_file.name.endswith(".pdf"):
                with open("temp.pdf", "wb") as f:
                    f.write(uploaded_file.read())
                references = extract_references_from_pdf("temp.pdf")
            elif uploaded_file.name.endswith(".docx"):
                with open("temp.docx", "wb") as f:
                    f.write(uploaded_file.read())
                references = extract_references_from_docx("temp.docx")
            else:
                references = uploaded_file.read().decode().splitlines()
        else:
            references = manual_refs.splitlines()

        results = []
        for ref in references:
            if ref.strip():
                st.markdown(f"### 🔍 Checking: {ref}")
                with st.spinner("Verifying..."):
                    result = verify_reference(ref.strip(), style)
                    results.append(result)
                    st.write(result)

        df = pd.DataFrame(results)

        # Download buttons
        csv_file = "verified_references.csv"
        df.to_csv(csv_file, index=False)
        with open(csv_file, "rb") as f:
            st.download_button("⬇️ Download CSV", f, file_name=csv_file, mime="text/csv")

        docx_file = "verified_references.docx"
        export_to_docx(df, docx_file)
        with open(docx_file, "rb") as f:
            st.download_button("⬇️ Download DOCX", f, file_name=docx_file)

        pdf_file = "verified_references.pdf"
        export_to_pdf(df, pdf_file)
        with open(pdf_file, "rb") as f:
            st.download_button("⬇️ Download PDF", f, file_name=pdf_file)

# Entry point
def main():
    import sys
    if len(sys.argv) == 1:
        streamlit_interface()
    else:
        file_path = sys.argv[1]
        style = sys.argv[2] if len(sys.argv) > 2 else "APA"
        if file_path.endswith(".pdf"):
            refs = extract_references_from_pdf(file_path)
        elif file_path.endswith(".docx"):
            refs = extract_references_from_docx(file_path)
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                refs = [line.strip() for line in f if line.strip()]

        for ref in refs:
            print(verify_reference(ref, style))

if __name__ == "__main__":
    main()
